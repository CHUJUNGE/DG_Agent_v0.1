from __future__ import annotations

import argparse
import json
import os
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

ROOT = Path(__file__).resolve().parents[1]
DEPS_DIR = ROOT / ".deps"
if DEPS_DIR.exists():
    sys.path.insert(0, str(DEPS_DIR))

import pdfplumber
from docx import Document

sys.path.insert(0, str(ROOT / "src"))

from prompt import ProjectFile, ProjectInfo, build_generation_prompt, compact_text


OUT_DIR = ROOT / "tests" / "prompt_tests"
DEFAULT_CASE_ROOT = ROOT / "case_data"
KNOWN_CASE_INFO: dict[str, ProjectInfo] = {
    "case_001": ProjectInfo(category="口香糖", brand="益达 / 绿箭", target_audience="18-40 Urban Striver"),
    "case_002": ProjectInfo(category="巧克力 / 零食", brand="德芙 / Mars", target_audience="巧克力目标消费者"),
}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def normalize_case_id(case_id: str) -> str:
    case_id = case_id.strip().lower()
    if not case_id:
        raise ValueError("case id cannot be empty")
    return case_id if case_id.lower().startswith("case_") else f"case_{case_id}"


def case_dir_for_id(case_root: Path, case_id: str) -> Path:
    suffix = case_id.split("_", 1)[1] if "_" in case_id else case_id
    candidates = [
        case_root / f"Case_{suffix}",
        case_root / case_id,
        case_root / case_id.upper(),
        case_root / case_id.capitalize(),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def discover_case_ids(case_root: Path) -> list[str]:
    if not case_root.exists():
        return []
    case_ids: list[str] = []
    for path in sorted(case_root.iterdir(), key=lambda p: p.name):
        if not path.is_dir():
            continue
        match = re.match(r"case[_-]?(.+)", path.name, flags=re.IGNORECASE)
        if match:
            case_ids.append(normalize_case_id(match.group(1)))
    return case_ids


def infer_file_type(path: Path) -> str:
    name = path.name.lower()
    original_name = path.name
    if "brief" in name:
        return "brief"
    if "proposal" in name:
        return "proposal"
    if "客户内部资料" in original_name or "internal" in name or "desk research" in name:
        return "client_material"
    if "final digital diary" in name or "final dg" in name or "最终" in original_name or "标答" in original_name:
        return "final_reference_excluded"
    return "other"


def project_info_for_case(case_id: str, case_dir: Path) -> ProjectInfo:
    info = KNOWN_CASE_INFO.get(case_id, ProjectInfo())
    metadata_path = case_dir / "case_info.json"
    default_notes = f"请基于 {case_id} 的输入材料生成题目设计；不要参考 Final Digital Diary DG / 最终标答。"

    if not metadata_path.exists():
        if not info.extra_notes:
            info.extra_notes = default_notes
        return info

    data = json.loads(metadata_path.read_text(encoding="utf-8-sig"))
    return ProjectInfo(
        category=data.get("category", info.category),
        brand=data.get("brand", info.brand),
        target_audience=data.get("target_audience", info.target_audience),
        has_idi=data.get("has_idi", info.has_idi),
        output_preference=data.get("output_preference", info.output_preference),
        extra_notes=data.get("extra_notes", info.extra_notes or default_notes),
    )


def extract_docx(path: Path, max_chars: int = 16000) -> str:
    doc = Document(str(path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        text = clean(para.text)
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return compact_text("\n".join(chunks), max_chars)


def extract_pdf(path: Path, max_pages: int = 8, max_chars: int = 18000) -> str:
    chunks: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for index, page in enumerate(pdf.pages[:max_pages], start=1):
            text = clean(page.extract_text() or "")
            if text:
                chunks.append(f"[Page {index}] {text}")
        if len(pdf.pages) > max_pages:
            chunks.append(f"[Note] PDF has {len(pdf.pages)} pages; only first {max_pages} pages extracted for prompt smoke test.")
    return compact_text("\n\n".join(chunks), max_chars)


def extract_pptx(path: Path, max_slides: int = 12, max_chars: int = 18000) -> str:
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    chunks: list[str] = []
    with zipfile.ZipFile(path) as z:
        slides = sorted(
            [name for name in z.namelist() if re.match(r"ppt/slides/slide\d+\.xml", name)],
            key=lambda name: int(re.findall(r"\d+", name)[-1]),
        )
        for slide_index, slide_name in enumerate(slides[:max_slides], start=1):
            xml = ET.fromstring(z.read(slide_name))
            texts = [clean(t.text or "") for t in xml.findall(".//a:t", ns) if clean(t.text or "")]
            if texts:
                chunks.append(f"[Slide {slide_index}] " + " | ".join(texts))
        if len(slides) > max_slides:
            chunks.append(f"[Note] PPTX has {len(slides)} slides; only first {max_slides} slides extracted for prompt smoke test.")
    return compact_text("\n\n".join(chunks), max_chars)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".pptx":
        return extract_pptx(path)
    if suffix in {".txt", ".md"}:
        return compact_text(path.read_text(encoding="utf-8-sig", errors="ignore"), 18000)
    return ""


def load_case_inputs(case_dir: Path) -> list[ProjectFile]:
    if not case_dir.exists():
        raise FileNotFoundError(f"Case folder not found: {case_dir}")

    files: list[ProjectFile] = []
    for path in sorted(case_dir.iterdir(), key=lambda p: p.name):
        if not path.is_file():
            continue
        file_type = infer_file_type(path)
        # During generation tests, do not include the final questionnaire/DG as input.
        if file_type == "final_reference_excluded":
            continue
        if path.name == "case_info.json" or path.suffix.lower() not in {".docx", ".pdf", ".pptx", ".txt", ".md"}:
            continue
        text = extract_text(path)
        if text:
            files.append(ProjectFile(filename=path.name, file_type=file_type, text=text))
    return files


def write_prompt_preview(case_id: str, project_info: ProjectInfo, files: list[ProjectFile]) -> Path:
    bundle = build_generation_prompt(project_info=project_info, files=files)
    out_path = OUT_DIR / f"{case_id}_prompt_preview.md"
    OUT_DIR.mkdir(exist_ok=True)

    lines: list[str] = []
    lines.append(f"# Prompt Preview - {case_id}")
    lines.append("")
    lines.append(f"- Selected cases: {', '.join(bundle.selected_case_ids) or 'none'}")
    lines.append(f"- Input files: {len(files)}")
    for file in files:
        lines.append(f"  - {file.filename} ({file.file_type}, {len(file.text)} chars)")
    lines.append("")

    for index, message in enumerate(bundle.messages, start=1):
        lines.append("---")
        lines.append("")
        lines.append(f"## Message {index}: {message.role}")
        lines.append("")
        lines.append("```text")
        lines.append(message.content)
        lines.append("```")
        lines.append("")

    out_path.write_text("\n".join(lines), encoding="utf-8-sig")
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build prompt previews for one or more local cases.")
    parser.add_argument("--case", action="append", help="Case id, for example case_005 or 005. Can be passed multiple times.")
    parser.add_argument("--case-root", default=os.environ.get("DG_AGENT_CASE_ROOT", str(DEFAULT_CASE_ROOT)))
    args = parser.parse_args()

    case_root = Path(args.case_root)
    case_ids = [normalize_case_id(case_id) for case_id in args.case] if args.case else discover_case_ids(case_root)
    if not case_ids:
        case_ids = ["case_001", "case_002"]

    for case_id in case_ids:
        case_dir = case_dir_for_id(case_root, case_id)
        project_info = project_info_for_case(case_id, case_dir)
        files = load_case_inputs(case_dir)
        out_path = write_prompt_preview(case_id, project_info, files)
        print(f"{case_id}: {out_path}")


if __name__ == "__main__":
    main()
